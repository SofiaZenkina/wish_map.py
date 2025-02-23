import streamlit as st
import pandas as pd
from datetime import datetime
from io import BytesIO
from docx import Document
from fpdf import FPDF

# Цветовая схема и стили
BUTTON_COLOR = "#fae7b5"
HEADER_FONT = "Lletraferida"
INFO_COLOR = "#d0e6a5"  # Мягкий зеленый для уведомлений

CUSTOM_CSS = f"""
    <style>
        /* Вкладки при наведении и активные */
        .stTabs [role="tab"]:hover {{
            background-color: {BUTTON_COLOR} !important;
            color: black !important;
        }}

        .stTabs [role="tab"][aria-selected="true"] {{
            background-color: {BUTTON_COLOR} !important;
            color: black !important;
        }}

        /* Кнопки без цвета при нажатии */
        div.stButton > button {{
            background-color: transparent !important;
            color: black !important;
            border: 1px solid {BUTTON_COLOR};
            border-radius: 5px;
        }}

        /* Заголовки */
        h1, h2, h3, h4, h5, h6 {{
            font-family: '{HEADER_FONT}', sans-serif;
        }}

        /* Надпись слева сверху */
        .custom-header {{
            position: absolute;
            top: 10px;
            left: 20px;
            font-family: '{HEADER_FONT}', sans-serif;
            font-style: italic;
            font-size: 16px;
            color: {BUTTON_COLOR};
        }}

        /* Уведомления */
        .stAlert {{
            background-color: {INFO_COLOR} !important;
            color: black !important;
        }}
    </style>
"""

# Добавление пользовательских стилей и надписи сверху
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)
st.markdown('<div class="custom-header">ваши цели ближе, чем вам кажется...</div>', unsafe_allow_html=True)

# Инициализация состояния для хранения желаний
if 'wishes' not in st.session_state:
    st.session_state['wishes'] = []

# Список категорий и месяцев
categories = ["Внешность", "Отношения", "Здоровье", "Финансы", "Карьера", "Саморазвитие", "Путешествия", "Хобби", "Другое"]
months = ["Январь", "Февраль", "Март", "Апрель", "Май", "Июнь", "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"]

# Создание вкладок
tabs = st.tabs(["Добавить желание", "Мои желания", "Скачать таблицу"])

# Вкладка для добавления желания
with tabs[0]:
    st.header("Добавить желание")
    category = st.selectbox("Выберите категорию желания:", categories)
    month = st.selectbox("Выберите месяц для реализации:", months)
    description = st.text_input("Опишите вашу цель или желание:")

    if st.button("Добавить желание"):
        if description.strip():
            new_wish = {
                "Желание": description.strip(),
                "Категория": category,
                "Месяц": month
            }
            st.session_state['wishes'].append(new_wish)
            st.success("Желание успешно добавлено!")
        else:
            st.warning("Пожалуйста, введите описание желания.")

# Вкладка для просмотра желаний
with tabs[1]:
    st.header("Мои желания")
    if st.session_state['wishes']:
        wishes_df = pd.DataFrame(st.session_state['wishes'])
        st.dataframe(wishes_df, use_container_width=True)
    else:
        st.info("Пока нет добавленных желаний. Добавьте их на вкладке 'Добавить желание'.")

# Вкладка для скачивания таблицы
with tabs[2]:
    st.header("Скачать таблицу")
    if st.session_state['wishes']:
        wishes_df = pd.DataFrame(st.session_state['wishes'])

        # Скачивание в формате Excel
        excel_buffer = BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
            wishes_df.to_excel(writer, index=False, sheet_name="Желания")
        excel_buffer.seek(0)
        st.download_button(
            label="Скачать в Excel",
            data=excel_buffer,
            file_name="мои_желания.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        # Скачивание в формате PDF
        pdf_buffer = BytesIO()
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Мои желания", ln=True, align='C')

        col_width = pdf.w / 4
        row_height = pdf.font_size * 1.5

        # Заголовки таблицы
        for header in wishes_df.columns:
            pdf.cell(col_width, row_height, header, border=1, align='C')
        pdf.ln(row_height)

        # Данные таблицы
        for _, row in wishes_df.iterrows():
            for item in row:
                pdf.cell(col_width, row_height, str(item), border=1, align='C')
            pdf.ln(row_height)

        pdf.output(pdf_buffer)
        pdf_buffer.seek(0)

        st.download_button(
            label="Скачать в PDF",
            data=pdf_buffer,
            file_name="мои_желания.pdf",
            mime="application/pdf"
        )

        # Скачивание в формате Word
        word_buffer = BytesIO()
        doc = Document()
        doc.add_heading("Мои желания", level=1)
        table = doc.add_table(rows=1, cols=len(wishes_df.columns))
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(wishes_df.columns):
            hdr_cells[i].text = column

        for _, row in wishes_df.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)

        doc.save(word_buffer)
        word_buffer.seek(0)

        st.download_button(
            label="Скачать в Word",
            data=word_buffer,
            file_name="мои_желания.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.info("Пока нет добавленных желаний для скачивания.")
